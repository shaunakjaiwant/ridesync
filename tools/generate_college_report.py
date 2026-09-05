import os
import sys
import subprocess
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()

    # Set page margins (0.8 in top/bottom, 0.85 in left/right)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Base Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = RGBColor(30, 41, 59) # Slate 800

    # Title Header Banner
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("RIDESYNC: AN INTELLIGENT CAMPUS MOBILITY, ROUTE MATCHING AND DRIVER DISPATCH PLATFORM")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(19)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(37, 99, 235) # Royal Blue

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(6)
    run_sub = p_sub.add_run("Comprehensive Technical & Architectural Project Report for Academic Evaluation")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11.5)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(71, 85, 105)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(14)
    r_meta = p_meta.add_run("System Architecture | Geodesic Route Fitting | Pessimistic Locking | Security & Verification | Mobile Responsive Design System")
    r_meta.font.size = Pt(9)
    r_meta.font.bold = True
    r_meta.font.color.rgb = RGBColor(15, 23, 42)

    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(30, 58, 138) # Dark Navy
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(11.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(29, 78, 216)
        return h

    def add_callout(text, bg_hex="F1F5F9"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, bg_hex)
        set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(9.5)
        r.font.italic = True
        r.font.color.rgb = RGBColor(30, 41, 59)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 1. Executive Summary
    add_heading_1("1. Executive Summary & Abstract")
    doc.add_paragraph(
        "RideSync is an end-to-end, high-performance web-based campus mobility and intelligent driver dispatch system. "
        "Engineered specifically to solve last-mile commute challenges across university ecosystems, RideSync provides "
        "a unified workspace for student riders, verified campus drivers, and system administrators. "
        "By integrating geodesic route-fit matching algorithms, real-time fallback dispatching, persistent session state management, "
        "and 1-tap emergency SOS notifications, RideSync guarantees campus safety, strict concurrency control, and operational efficiency."
    )
    add_callout(
        "Key Project Highlights & Achievements:\n"
        "• Multi-Role Mobile Responsive Architecture: Unified Rider, Driver, and Admin portals with sticky glassmorphic navigation, safe-area inset compliance, and responsive bottom dock.\n"
        "• Algorithmic Route Matching: Haversine geodesic distance calculations, vector direction alignment, and time window fit scoring.\n"
        "• High Concurrency & Lock Protection: Row-level SELECT ... FOR UPDATE database transaction locks preventing double-claim race conditions.\n"
        "• Enhanced Security & Persistent Sessions: 30-day sliding session windows with instant middleware revocation for suspended user accounts.\n"
        "• Automated AI Driver KYC Verification: Python FastAPI microservice utilizing OpenCV feature matching and document OCR for automated driver onboarding.\n"
        "• Sub-50ms API Performance & 100% Quality Pass: OPcache TTFB optimizations, session principal caching, and 65 automated unit tests passing in PHPUnit with zero regression vulnerabilities."
    )

    # 2. System Architecture
    add_heading_1("2. Technology Stack & System Architecture")
    doc.add_paragraph(
        "RideSync is engineered using a decoupled, modular Model-View-Controller (MVC) architecture in PHP 8.2 and MySQL/MariaDB 8.4, "
        "complemented by a responsive Vanilla JavaScript frontend with centralized CSS design tokens, a Node.js WebSocket gateway for real-time telemetry, "
        "and an external Python AI KYC document verification service."
    )

    t_stack = doc.add_table(rows=1, cols=3)
    t_stack.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t_stack.rows[0].cells
    hdr[0].text = "Layer / Component"
    hdr[1].text = "Technology / Framework"
    hdr[2].text = "Technical Purpose & Responsibility"
    for cell in hdr:
        set_cell_background(cell, "1E293B")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(255, 255, 255)

    stack_data = [
        ("Backend Web Core", "PHP 8.2 (Strict Types, Mysqli, OPcache)", "Request handling, auth middleware, session principal caching, CSRF token verification, action routing"),
        ("Database Layer", "MySQL 8.4 / MariaDB 10.4+ (Aiven SSL CA)", "ACID relational persistence, spatial indexes, FOR UPDATE pessimistic locks, automated schema migrations"),
        ("Frontend & Design System", "Vanilla JS (ES6+), Glassmorphic CSS", "Design tokens scale, responsive glassmorphic cards, notification toasts, live tracking, mobile bottom dock"),
        ("Geolocation & Maps", "OSRM & Leaflet / OpenStreetMap", "Geodesic coordinate routing, pickup ETA calculation, route line polyline rendering"),
        ("Realtime Telemetry", "Node.js 20+ WebSocket Gateway", "HMAC SHA-256 signed query tokens, low-latency SOS emergency alerts, driver live location broadcast"),
        ("AI Driver Verification", "Python 3.11/3.13 FastAPI + OpenCV", "Automated driver KYC document OCR, facial feature matching, and background verification service"),
        ("Cloud Infrastructure", "Render Docker Blueprints & Vercel Edge", "Vercel edge reverse proxying, Render Web Service containerization, S3 document storage abstraction"),
        ("Quality & Testing", "PHPUnit 11 & K6 Load Testing", "65 automated unit tests, negative security tests, session security checks, sub-50ms performance load profiling")
    ]

    for cat, tech, desc in stack_data:
        row = t_stack.add_row().cells
        row[0].text = cat
        row[1].text = tech
        row[2].text = desc
        for cell in row:
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 3. Algorithms & Math
    add_heading_1("3. Core Mathematical Models & Algorithms")
    
    add_heading_2("3.1 Geodesic Route Distance (Haversine Formula)")
    doc.add_paragraph(
        "To compute precise spatial distances between pickup coordinates (lat1, lon1) and destination coordinates (lat2, lon2) on Earth, "
        "RideSync implements the spherical Haversine formula:"
    )
    add_callout(
        "a = sin²(Δlat / 2) + cos(lat1) * cos(lat2) * sin²(Δlon / 2)\n"
        "c = 2 * atan2(√a, √(1 − a))\n"
        "Distance (d) = R * c   (where Earth Radius R = 6,371 km)"
    )

    add_heading_2("3.2 Route Fit & Compatibility Scoring Equation")
    doc.add_paragraph(
        "When a rider searches for rides, RideSync evaluates every candidate route using a multi-factor fitness function:"
    )
    add_callout(
        "Fit Score = (W_dist * S_dist) + (W_time * S_time) + (W_seats * S_seats)\n\n"
        "Where:\n"
        "• S_dist = max(0, 1 - (Pickup_Deviation_km / Max_Allowed_Radius))\n"
        "• S_time = max(0, 1 - (|Requested_Time - Offered_Time|_mins / 60))\n"
        "• S_seats = (Available_Seats >= Requested_Seats) ? 1.0 : 0.0\n"
        "• Weights: W_dist = 0.50, W_time = 0.35, W_seats = 0.15"
    )

    add_heading_2("3.3 Concurrency & Pessimistic Lock Prevention Algorithm")
    doc.add_paragraph(
        "To eliminate race conditions when multiple drivers attempt to claim the same rider request simultaneously, "
        "RideSync enforces database-level transactions with SELECT ... FOR UPDATE locks:"
    )
    add_callout(
        "1. Start DB Transaction: mysqli_begin_transaction($conn);\n"
        "2. Acquire Exclusive Lock: SELECT status FROM driver_requests WHERE id = ? FOR UPDATE;\n"
        "3. Validate Status: If status != 'pending', ROLLBACK transaction and return HTTP 409 conflict error.\n"
        "4. Mutate State: UPDATE driver_requests SET status = 'accepted', driver_id = ? WHERE id = ?;\n"
        "5. Commit Transaction: mysqli_commit($conn);"
    )

    add_heading_2("3.4 Persistent 30-Day Session & Instant Revocation Algorithm")
    doc.add_paragraph(
        "RideSync implements a non-intrusive persistent session model that keeps authenticated users logged in for 30 days "
        "across browser restarts, combined with immediate security enforcement:"
    )
    add_callout(
        "• Session Cookie Lifetime: 2,592,000 seconds (30 Days)\n"
        "• Garbage Collection Max Lifetime: ini_set('session.gc_maxlifetime', '2592000');\n"
        "• Sliding Session Expiry: Inactivity timeout refreshed on every request up to 30 days.\n"
        "• Instant Revocation Middleware: ridesync_validate_session_principal checks user status in DB on every request; if status == 'suspended', session is terminated immediately and cookies destroyed."
    )

    # 4. Mobile Responsive Design System
    add_heading_1("4. Mobile Responsive Design System & UI Architecture")
    doc.add_paragraph(
        "RideSync features a modern, mobile-first design system built using Vanilla CSS tokens, glassmorphic styling, "
        "and progressive enhancement without heavy external UI dependencies."
    )
    add_callout(
        "Design System Specifications:\n"
        "• Single-Row Sticky Glassmorphic Navbar: Compact top bar housing logo mark, page header, unread alert bell badge, and user profile avatar dropdown.\n"
        "• Mobile Bottom Navigation Dock: Fixed bottom tab bar providing 1-tap switching between Home, Search/Post, Driver Roster, Reports, and Settings.\n"
        "• Responsive Table-to-Card Transformations: Complex tabular data automatically transforms into stacked glassmorphic cards on mobile viewports (<768px).\n"
        "• Ergonomic Touch Targets & Protection: Minimum 44px x 44px interactive tap targets, safe-area-inset padding for notched devices, and disabled auto-zoom (`touch-action: manipulation`)."
    )

    # 5. Database Schema
    add_heading_1("5. Database Schema & Data Dictionary")
    doc.add_paragraph(
        "The RideSync database contains 12 core relational tables configured with foreign key constraints, "
        "composite spatial/telemetry indexes, and full ACID compliance."
    )

    t_db = doc.add_table(rows=1, cols=4)
    t_db.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t_db.rows[0].cells
    hdr[0].text = "Table Name"
    hdr[1].text = "Primary Keys / Indexes"
    hdr[2].text = "Key Attributes"
    hdr[3].text = "System Purpose"
    for cell in hdr:
        set_cell_background(cell, "1E293B")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(255, 255, 255)

    db_data = [
        ("users", "PRIMARY(id), UNIQUE(email)", "name, email, password, status, role", "Rider account credentials & instant suspension status"),
        ("driver_accounts", "PRIMARY(id), UNIQUE(email, phone)", "name, phone, status, availability", "Driver credentials, approval state & online availability status"),
        ("driver_documents", "PRIMARY(id), idx_driver_doc", "driver_id, document_type, status", "KYC documents (License, Aadhaar, PAN, RC) & AI review state"),
        ("rides", "PRIMARY(id), idx_status_date", "user_id, origin, destination, seats", "Posted shared trips and pickup/destination criteria"),
        ("driver_requests", "PRIMARY(id), idx_ride_driver", "ride_id, driver_id, status, fare", "Direct rider-to-driver dispatch requests with pessimistic locks"),
        ("sos_alerts", "PRIMARY(id), idx_status_id", "ride_id, triggered_by_type, lat, lng", "Active emergency SOS alerts & location coordinates"),
        ("user_emergency_contacts", "PRIMARY(id), idx_user_role", "user_id, user_type, name, phone", "Personal emergency contacts notified on SOS trigger"),
        ("realtime_events", "PRIMARY(id), UNIQUE(idempotency)", "event_type, audience, payload_json", "Idempotent event log for WebSocket push notifications")
    ]

    for tbl_name, pkey, attrs, purpose in db_data:
        row = t_db.add_row().cells
        row[0].text = tbl_name
        row[1].text = pkey
        row[2].text = attrs
        row[3].text = purpose
        for cell in row:
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 6. API Endpoints
    add_heading_1("6. API Specifications & Action Endpoints")
    doc.add_paragraph(
        "RideSync exposes modular action handlers and RESTful JSON endpoints for real-time interactions:"
    )

    t_api = doc.add_table(rows=1, cols=4)
    t_api.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t_api.rows[0].cells
    hdr[0].text = "Endpoint / Action Route"
    hdr[1].text = "Method"
    hdr[2].text = "Parameters / Payload"
    hdr[3].text = "Response / Effect"
    for cell in hdr:
        set_cell_background(cell, "1E293B")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(255, 255, 255)

    api_data = [
        ("/api/v1/user_notifications_check.php", "GET", "since_id (int)", "JSON unread count and latest alert toast payload"),
        ("/actions/login_action.php", "POST", "email, password, csrf_token", "Authenticates user, checks suspension status, sets 30-day session"),
        ("/actions/match_action.php", "POST", "ride_id, action_type", "Creates join request or updates match state with transaction locks"),
        ("/actions/driver_request_action.php", "POST", "action_type, ride_id, fare", "Executes FOR UPDATE locked driver request claim"),
        ("/actions/emergency_contact_action.php", "POST", "action_type, name, phone", "Adds or removes user emergency safety contacts"),
        ("/actions/sos_action.php", "POST", "ride_id, latitude, longitude", "Triggers SOS alert, broadcasts WebSocket event & notifies safety contacts")
    ]

    for ep, m, params, resp in api_data:
        row = t_api.add_row().cells
        row[0].text = ep
        row[1].text = m
        row[2].text = params
        row[3].text = resp
        for cell in row:
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 7. Cloud Infrastructure & Security Architecture
    add_heading_1("7. Security Architecture & Cloud Deployment")
    doc.add_paragraph(
        "RideSync incorporates enterprise-grade security mechanisms across data persistence, session management, and deployment topology:"
    )
    add_callout(
        "Security & DevOps Specifications:\n"
        "• Sensitive Document Encryption: Driver KYC documents stored locally or in S3 are encrypted at rest using AES-256-CBC via RIDESYNC_DOCUMENT_ENCRYPTION_KEY.\n"
        "• CSRF & SQL Injection Protection: All state-changing POST endpoints enforce anti-CSRF token verification; database interactions use prepared mysqli statements.\n"
        "• Content Security Policy (CSP): Dynamic CSP header supporting local/LAN IP WebSocket connections (`connect-src ws://127.0.0.1:8081`).\n"
        "• Monorepo Cloud Topology: Vercel Edge reverse proxy (`vercel.json`) routing traffic to Render Docker Web Service (`render.yaml`) connected to Managed Aiven MySQL via SSL CA certificate (`ca.pem`)."
    )

    # 8. Quality Assurance & Performance Metrics
    add_heading_1("8. Quality Assurance, Testing & Performance Metrics")
    doc.add_paragraph(
        "RideSync enforces automated quality gates prior to any production deployment:"
    )
    add_callout(
        "• PHPUnit Test Suite: 65 automated unit tests covering authentication security, OTP validation, pessimistic locks, session security, and emergency contacts.\n"
        "• TTFB Latency Optimization: OPcache bytecode caching, session principal caching, early session unlock, and Gzip compression delivering ultra-low page rendering latencies.\n"
        "• High Concurrency Stress Testing: K6 load testing validating sub-50ms API response times under high concurrent user load."
    )

    # Save DOCX file
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    docs_dir = os.path.join(project_root, "docs")
    docx_path = os.path.join(docs_dir, "RideSync_Comprehensive_College_Project_Report.docx")

    os.makedirs(docs_dir, exist_ok=True)
    doc.save(docx_path)
    print(f"DOCX Report created successfully at:\n  {docx_path}")

    # Export to PDF via convert_docx_to_pdf.ps1
    pdf_path = os.path.join(docs_dir, "RideSync_Report.pdf")
    ps_script = os.path.join(script_dir, "convert_docx_to_pdf.ps1")
    
    print("\nExporting PDF report...")
    cmd = ["powershell", "-ExecutionPolicy", "Bypass", "-File", ps_script, "-DocxPath", docx_path, "-PdfPath", pdf_path]
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode == 0:
        print(f"PDF Report generated successfully at: {pdf_path}")
    else:
        print(f"Warning: PDF conversion finished with output:\n{res.stdout}\n{res.stderr}")

if __name__ == "__main__":
    create_report()
